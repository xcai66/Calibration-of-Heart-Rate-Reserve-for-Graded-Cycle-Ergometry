#!/usr/bin/env python3
"""Build submission-facing DOCX files from the final Markdown sources.

Design system: narrative_proposal preset with a named scientific-manuscript
override (A4, Times New Roman/Hiragino Sans GB, restrained black hierarchy, 1.5-line
body text, fixed-width landscape main tables). First-page pattern:
editorial_cover adapted to a conventional scientific title page.
"""

from __future__ import annotations

import argparse
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


NAVY = "173A53"
BLUE = "1E5A8A"
MUTED = "5F6B73"
LIGHT = "E8EEF5"
GRID = "B8C3CC"
BLACK = "000000"

FIGURE_LEGENDS = {
    1: "Study design and development-only sport selection. Development, temporal-holdout, and external-validation data were separated before final testing. Cycling was selected using development performance, sample size, and independent-data availability.",
    2: "Locked CycHRR-T function and validation performance. The fixed transfer was evaluated against raw HRR and development-fitted linear comparators in temporal and external data. VO2R denotes oxygen-uptake reserve; MAE denotes mean absolute error.",
    3: "Paired effects, strong-comparator audit, and HR-anchor sensitivity. Negative MAE differences favor CycHRR-T. Anchor perturbations show that maximal-HR specification can change the direction of model contrast.",
}

FIGURE_LEGENDS_CN = {
    1: "研究设计与仅基于开发数据的项目筛选。开发集、时间外推保留集和外部验证集在最终检验前保持隔离。自行车根据开发表现、样本量和独立数据可得性被选中。",
    2: "锁定的CycHRR-T函数及验证表现。固定转换函数在时间外推和外部数据中与原始HRR及开发数据拟合的线性对照进行比较。VO2R表示摄氧量储备，MAE表示平均绝对误差。",
    3: "配对效应、强线性对照核查和HR锚点敏感性。负MAE差值有利于CycHRR-T。锚点扰动表明，最大HR设定可改变模型差值方向。",
}


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=GRID, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, font_name, east_asia, size=None, bold=None, italic=None, color=BLACK):
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, fld_sep, text, fld_end])


def configure_section(section, landscape=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.9)
        section.right_margin = Cm(1.9)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def configure_styles(doc, chinese=False):
    latin = "Hiragino Sans GB" if chinese else "Times New Roman"
    east = latin
    normal = doc.styles["Normal"]
    normal.font.name = latin
    normal._element.rPr.rFonts.set(qn("w:ascii"), latin)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5

    specs = {
        "Title": (17, NAVY, 0, 10),
        "Heading 1": (15, BLACK, 16, 8),
        "Heading 2": (12.5, BLACK, 12, 6),
        "Heading 3": (11.5, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = latin
        style._element.rPr.rFonts.set(qn("w:ascii"), latin)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east)
        style.font.size = Pt(size)
        style.font.bold = name != "Title"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.15

    for name in ("Caption",):
        style = doc.styles[name]
        style.font.name = latin
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east)
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor.from_string(MUTED)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.05


def set_headers_footers(section, left_text, chinese=False, first=False):
    latin = "Hiragino Sans GB" if chinese else "Times New Roman"
    east = latin
    section.different_first_page_header_footer = first
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(left_text)
    set_run_font(r, latin, east, 8.5, color=MUTED)
    if first:
        fp = section.first_page_header.paragraphs[0]
        fp.text = ""
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    label = fp.add_run("Page " if not chinese else "第 ")
    set_run_font(label, latin, east, 8.5, color=MUTED)
    add_page_field(fp)
    if chinese:
        tail = fp.add_run(" 页")
        set_run_font(tail, latin, east, 8.5, color=MUTED)


def add_inline(paragraph, text, chinese=False, base_size=11):
    latin = "Hiragino Sans GB" if chinese else "Times New Roman"
    east = latin
    # Bold and italic markdown are sufficient for the final sources.
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*)")
    for part in pattern.split(text):
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        italic = part.startswith("*") and part.endswith("*") and not bold
        clean = part[2:-2] if bold else part[1:-1] if italic else part
        run = paragraph.add_run(clean)
        set_run_font(run, latin, east, base_size, bold=bold, italic=italic)


def add_body_paragraph(doc, text, chinese=False, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.widow_control = True
    add_inline(p, text, chinese=chinese)
    return p


def add_formula(doc, text, chinese=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    formula_font = "Hiragino Sans GB" if chinese else "Cambria Math"
    set_run_font(r, formula_font, formula_font, 11.5, italic=True)
    return p


def set_table_geometry(table, widths_inches):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_dxa = int(round(sum(widths_inches) * 1440))
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_inches[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(int(round(width * 1440))))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def table_widths(ncols, total=10.2):
    patterns = {
        6: [1.45, 1.55, 1.75, 0.85, 1.50, 1.60],
        9: [1.35, 1.15, 1.05, 0.50, 0.85, 0.85, 1.80, 0.80, 0.65],
    }
    base = patterns.get(ncols, [1.0] * ncols)
    scale = total / sum(base)
    return [value * scale for value in base]


def add_markdown_table(doc, lines, chinese=False):
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-+:?", cell or "-") for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    set_table_geometry(table, table_widths(ncols))
    set_table_borders(table)
    latin = "Hiragino Sans GB" if chinese else "Times New Roman"
    east = latin
    for i, values in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = values[j] if j < len(values) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(text)
            set_run_font(r, latin, east, 7.7, bold=i == 0)
            if i == 0:
                set_cell_shading(cell, LIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, number, path, chinese=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(str(path), width=Inches(6.05))
    caption = doc.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label = f"Figure {number}. " if not chinese else f"图{number} "
    r = caption.add_run(label)
    figure_font = "Hiragino Sans GB" if chinese else "Times New Roman"
    set_run_font(r, figure_font, figure_font, 9, bold=True, color=MUTED)
    legend = FIGURE_LEGENDS_CN[number] if chinese else FIGURE_LEGENDS[number]
    rr = caption.add_run(legend)
    set_run_font(rr, figure_font, figure_font, 9, color=MUTED)


def build_manuscript(source, output, figures_dir, chinese=False):
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_section(doc.sections[0], landscape=False)
    configure_styles(doc, chinese=chinese)
    running = "CycHRR-T | Scientific manuscript" if not chinese else "CycHRR-T | 科学论文稿"
    set_headers_footers(doc.sections[0], running, chinese=chinese, first=True)

    in_title_block = True
    abstract_started = False
    skip_legends = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|-") or stripped.startswith("|:"):
            i += 1
            continue

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_markdown_table(doc, block, chinese=chinese)
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(14)
            r = p.add_run(stripped[2:].strip())
            title_font = "Hiragino Sans GB" if chinese else "Times New Roman"
            set_run_font(r, title_font, title_font, 17, bold=True, color=NAVY)
            i += 1
            continue

        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            if heading in ("Figure legends", "图注"):
                skip_legends = True
                i += 1
                continue
            if heading in ("Main tables", "主要表格"):
                skip_legends = False
                section = doc.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=True)
                set_headers_footers(section, running, chinese=chinese, first=False)
                p = doc.add_paragraph(heading, style="Heading 1")
                p.paragraph_format.space_before = Pt(0)
                i += 1
                continue
            if skip_legends:
                i += 1
                continue
            if heading in ("Abstract", "摘要") and not abstract_started:
                doc.add_page_break()
                abstract_started = True
                in_title_block = False
            doc.add_paragraph(heading, style="Heading 1")
            i += 1
            continue

        if skip_legends:
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:].strip(), style="Heading 2")
            i += 1
            continue

        if stripped.startswith("#### "):
            doc.add_paragraph(stripped[5:].strip(), style="Heading 3")
            i += 1
            continue

        figure_markdown = re.match(r"!\[(?:Figure|图)\s*(\d)[^\]]*\]\(([^)]+)\)", stripped)
        if figure_markdown:
            number = int(figure_markdown.group(1))
            linked_path = (source.parent / figure_markdown.group(2)).resolve()
            if linked_path.exists():
                figure_path = linked_path
            else:
                candidates = sorted(figures_dir.glob(f"Figure_{number}_*.png"))
                if not candidates:
                    raise FileNotFoundError(f"Figure {number} PNG not found")
                figure_path = candidates[0]
            add_figure(doc, number, figure_path, chinese=chinese)
            i += 1
            continue

        if in_title_block and stripped.startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, stripped, chinese=chinese, base_size=10.5)
            i += 1
            continue

        if stripped.startswith("**Table ") or stripped.startswith("**表"):
            p = doc.add_paragraph(style="Caption")
            p.paragraph_format.keep_with_next = True
            add_inline(p, stripped, chinese=chinese, base_size=9)
            i += 1
            continue

        if re.match(r"^(g\(h\)|h =|v =|p =)", stripped):
            add_formula(doc, stripped, chinese=chinese)
            i += 1
            continue

        if re.match(r"^\d+\. ", stripped) and stripped.count(".") == 1:
            p = add_body_paragraph(doc, stripped, chinese=chinese)
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            i += 1
            continue

        if re.match(r"^\d+\. ", stripped):
            p = add_body_paragraph(doc, stripped, chinese=chinese)
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(3)
            i += 1
            continue

        if stripped.startswith("- "):
            p = add_body_paragraph(doc, "• " + stripped[2:], chinese=chinese)
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            i += 1
            continue

        p = add_body_paragraph(doc, stripped, chinese=chinese)
        if stripped.startswith("**Keywords:") or stripped.startswith("**关键词："):
            p.paragraph_format.space_after = Pt(10)
        i += 1

    doc.core_properties.title = lines[0].lstrip("# ")
    doc.core_properties.author = "BoTao Cai"
    doc.core_properties.subject = "Cycling-specific heart-rate-reserve transfer function"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_short(source, output, title, chinese=False, letter=False):
    doc = Document()
    configure_section(doc.sections[0], landscape=False)
    configure_styles(doc, chinese=chinese)
    if title == "Title page and declarations":
        section = doc.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        normal = doc.styles["Normal"]
        normal.font.size = Pt(10)
        normal.paragraph_format.line_spacing = 1.15
        normal.paragraph_format.space_after = Pt(4)
        for heading_name in ("Heading 1", "Heading 2"):
            doc.styles[heading_name].paragraph_format.space_before = Pt(6)
            doc.styles[heading_name].paragraph_format.space_after = Pt(3)
    if letter:
        section = doc.sections[0]
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        normal = doc.styles["Normal"]
        normal.font.size = Pt(10)
        normal.paragraph_format.line_spacing = 1.05
        normal.paragraph_format.space_after = Pt(4)
        doc.styles["Title"].font.size = Pt(15)
        doc.styles["Title"].paragraph_format.space_after = Pt(6)
    set_headers_footers(doc.sections[0], title, chinese=chinese, first=True)
    lines = source.read_text(encoding="utf-8").splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if not letter else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, stripped[2:], chinese=chinese, base_size=17)
        elif stripped.startswith("## "):
            doc.add_paragraph(stripped[3:], style="Heading 1")
        elif stripped.startswith("- "):
            p = add_body_paragraph(doc, "• " + stripped[2:], chinese=chinese)
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.first_line_indent = Cm(-0.5)
        else:
            p = add_body_paragraph(doc, stripped, chinese=chinese)
            if letter:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.core_properties.title = title
    doc.core_properties.author = "BoTao Cai"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=Path, required=True)
    args = parser.parse_args()
    root = args.root.resolve()
    manuscripts = root / "05_manuscript"
    submission = root / "06_submission"
    figures = root / "04_results" / "figures"
    build_manuscript(manuscripts / "English_Manuscript.md", manuscripts / "English_Manuscript.docx", figures, False)
    build_manuscript(manuscripts / "Chinese_Manuscript.md", manuscripts / "Chinese_Manuscript.docx", figures, True)
    build_short(submission / "Title_Page_and_Declarations.md", submission / "Title_Page_and_Declarations.docx", "Title page and declarations")
    build_short(submission / "Cover_Letter_Draft.md", submission / "Cover_Letter_Draft.docx", "Cover letter", letter=True)
    build_short(submission / "Highlights.md", submission / "Highlights.docx", "Highlights")


if __name__ == "__main__":
    main()
